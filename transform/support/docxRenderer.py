# -*- coding: utf-8 -*-
#

import codecs
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor

import abstractRenderer


#
#   Renders to Word .docx
#

class Renderer(abstractRenderer.AbstractRenderer):
    
    def __init__(self, inputDir, outputDir, outputName, config):
        abstractRenderer.AbstractRenderer.__init__(self, inputDir, outputDir, outputName, config)
        self.document = Document()
        self.currentParagraph = self.document.add_paragraph()
        self.outputFilename = os.path.join(outputDir, outputName + '.docx')
        self.inputDir = inputDir
        self.f = codecs.open('/dev/null', 'w', 'utf_8_sig')
        self.currentChapter = ''
        self.currentBook = ''
        self.in_nd = False
        self.in_footnote = False
        # Old:
        # Flags
        self.printerState = {'li': False, 'd': False, 'm': False}
        self.smallCapSections = True  # Sometimes we don't want to do this, like for Psalms
        self.justDidLORD = False
        self.justDidNB = False
        self.doNB = False
        self.narrower = False
        self.doChapterOrVerse = ''
        self.smallcaps = False

    def render(self, order='normal'):
        self.loadUSFM(self.inputDir)
        self.run(order)
        self.document.save(self.outputFilename)        
    
    #
    #   Support
    #

    def startNarrower(self, n):
        s = '}' if self.narrower else '\n\\blank[medium] '
        self.narrower = True
        s = s + '\n\\noindentation \\Q{' + str(n) + '}{'
        self.doNB = True
        return s

    def stopNarrower(self):
        s = '}\n\\blank[medium] ' if self.narrower else ''
        self.narrower = False
        return s

    def escapeText(self, s):
        return s.replace('&', '\\&').replace('%', '\\%')
          
    def startLI(self):
        if self.printerState['li'] == False:
            self.printerState['li'] = True
            #return u'\startitemize \item '
            return r'\startexdent '
        else:
            #return u'\item '
            return r'\par '
        
    def stopLI(self):
        if self.printerState['li'] == False:
            return ''
        else:
            self.printerState['li'] = False
            #return u'\stopitemize'
            return r'\stopexdent '

    def startD(self):
        if self.printerState['d'] == False:
            self.printerState['d'] = True
        return '\par {\startalignment[middle] \em '

    def stopD(self):
        if self.printerState['d'] == False:
            return ''
        else:
            self.printerState['d'] = False
            return '\stopalignment }'
            
    def startM(self):
        r = self.stopD() + self.stopLI() + self.stopNarrower()
        self.printerState['m'] = True
        return r + r'\par {\startalignment[flushleft] '

    def stopM(self):
        if self.printerState['m'] == False:
            return ''
        else:
            self.printerState['m'] = False
            return '\stopalignment }'

    def newLine(self):
        s = '\n\par \n'
        if self.doNB:
            self.doNB = False
            self.justDidNB = True
            s = s + r'\noindentation '
        elif self.justDidNB:
            self.justDidNB = False
            s = s + r'\indentation '
        return s
                    
    #
    #   Used helpers for Docx
    #
    def clean(self, token):
        return token.getValue().replace('~', ' ')
        
    def newPara(self):
        self.currentParagraph = self.document.add_paragraph()

    #
    #   Tokens
    #

    def render_h(self, token):       self.document.add_page_break(); self.currentBook = token.getValue()
    def render_mt1(self, token):     self.document.add_heading(self.clean(token), level=0)
    def render_mt2(self, token):     self.document.add_heading(self.clean(token), level=1)
    def render_ms1(self, token):     self.document.add_heading(self.clean(token), level=2)
    def render_ms2(self, token):     self.document.add_heading(self.clean(token), level=3)
    def render_p(self, token):       self.newPara(); 
    def render_pi(self, token):      self.newPara(); self.currentParagraph.left_indent = Inches(1)

    def render_s1(self, token):      self.document.add_heading(self.clean(token), level=4)
    def render_s2(self, token):      self.document.add_heading(self.clean(token), level=5)

    def render_c(self, token):        self.currentChapter = token.getValue()
    def render_v(self, token):
        if token.getValue() == '1':
            run = self.currentParagraph.add_run(self.currentBook + ' ' + self.currentChapter + ' ')
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run = self.currentParagraph.add_run(token.getValue() + ' ')
            run.font.color.rgb = RGBColor(0, 140, 0)
        run.font.superscript = True
            
    def render_text(self, token):
        run = self.currentParagraph.add_run(self.clean(token) + ' ')
        if self.in_nd: 
            run.font.small_caps = True
        elif self.in_footnote:
            run.font.color.rgb = RGBColor(0, 0, 140)
        else:
            pass
        
    def render_q1(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0)
    def render_q2(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0); self.currentParagraph.add_run('\t')
    def render_q3(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0); self.currentParagraph.add_run('\t\t')
    def render_nb(self, token):     self.newPara(); self.currentParagraph.left_indent = Inches(0)
    def render_b(self, token):      self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0)

    def render_d(self, token):      self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0)

    def render_f_s(self, token):
        run = self.currentParagraph.add_run(' [[ ')
        run.font.color.rgb = RGBColor(0, 0, 140)
        self.in_footnote = True
    def render_f_e(self, token):
        run = self.currentParagraph.add_run(' ]] ')
        run.font.color.rgb = RGBColor(0, 0, 140)
        self.in_footnote = False
    
    def render_nd_s(self, token): pass
    def render_nd_e(self, token):
        run = self.currentParagraph.runs[-1]
        run.font.small_caps = True

    def render_em_s(self, token): pass
    def render_em_e(self, token):
        run = self.currentParagraph.runs[-1]
        run.italic = True

    # SELAH
    def render_qs_s(self, token): pass
    def render_qs_e(self, token):
        run = self.currentParagraph.runs[-1]
        run.italic = True

    def render_wj_s(self, token): pass
    def render_wj_e(self, token):
        run = self.currentParagraph.runs[-1]
        run.font.color.rgb = RGBColor(140, 0, 0)
    
    def render_is1(self, token):    self.render_s1(token)
    def render_ip(self, token):     self.render_p(token)
    def render_iot(self, token):    self.render_q1(token)
    def render_io1(self, token):    self.render_q2(token)
    
    def render_pb(self, token):     self.document.add_page_break()
    def render_m(self, token):      self.render_p(token)
    
    #
    #   Introductory codes
    #
    
    introTeXt = str(r""" """)
    closeTeXt = r""" """
